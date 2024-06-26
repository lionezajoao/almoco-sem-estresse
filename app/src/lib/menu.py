import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from app.src.lib.email_sender import EmailSender
from app.src.database.menu import MenuDatabase
from app.models.menu_models import NewMenuModel

class Menu(MenuDatabase):
    def __init__(self):
        super().__init__()
        self.email = EmailSender()

    @MenuDatabase.log
    def get_all_items(self):
        self.connect()
        data = self.get_menu()
        self.close()
        return data
    
    @MenuDatabase.log
    def get_item_by_type(self, type):
        self.connect()
        data = self.get_data_by_type(type)
        self.close()
        return [item[0] for item in data]
    
    @MenuDatabase.log
    def get_all_ingredients(self):
        self.connect()
        data = self.get_ingredients()
        self.close()
        return data
    
    @MenuDatabase.log
    def get_data_by_name(self, name):
        self.connect()
        data = self.get_item_by_name(name)
        self.close()
        return data
    
    @MenuDatabase.log
    def get_ingredients_from_items(self, name):
        self.connect()
        data = self.get_items_ingredients(name)
        self.close()
        return data
    
    @MenuDatabase.log
    def get_all_items_ingredients(self):
        return self.get_relational_menu()
    
    @MenuDatabase.log
    def get_ingredient_type(self, ingredient):
        self.connect()
        data = self.get_ingredient_type_by_name(ingredient)
        self.close()
        return data
    
    @MenuDatabase.log
    def insert_new_item(self, data):
        self.connect()
        if not self.check_if_item_exists(data.name):
            self.insert_item(data.name)

        for ingredient in data.ingredients:
            if not self.check_if_ingredient_exists(ingredient):
                self.insert_ingredient(ingredient)
            self.insert_item_ingredient(data.name, ingredient)

        self.close()
    
    @MenuDatabase.log
    def create_menu(self, data: NewMenuModel, token_data: dict):
        self.connect()
        relational_data = self.transform_dataset(self.get_all_items_ingredients())
        self.close()
        menu_data = {}
        ingredients_data = {}
        menu_name = f"Cardápio - { token_data.get('name') }"

        for item in data:
            menu_data[item.week_choice] = {}
            for day in item.data:
                main_dish_data = relational_data.get(day.main_dish)
                salad_data = relational_data.get(day.salad)

                menu_data[item.week_choice][day.weekday] = {
                    "main_dish": {
                        "dish": day.main_dish,
                        "ingredients": {
                            "protein": main_dish_data.get("proteina", []) if main_dish_data else [],
                            "hortifrutti": main_dish_data.get("hortifruti", []) if main_dish_data else [],
                            "cold_cuts": main_dish_data.get("frios", []) if main_dish_data else [],
                            "grocery": main_dish_data.get("mercearia", []) if main_dish_data else []
                        }
                    },
                    "salad": {
                        "dish": day.salad,
                        "ingredients": {
                            "protein": salad_data.get("proteina", []) if salad_data else [],
                            "hortifrutti": salad_data.get("hortifruti", []) if salad_data else [],
                            "cold_cuts": salad_data.get("frios", []) if salad_data else [],
                            "grocery": salad_data.get("mercearia", []) if salad_data else []
                        }
                    },
                    "accompaniment": []
                }
                for accompaniment in day.accompaniment:
                    accompaniment_data = relational_data.get(accompaniment)
                    menu_data[item.week_choice][day.weekday]["accompaniment"].append({
                        "dish": accompaniment,
                        "ingredients": {
                            "protein": accompaniment_data.get("proteina", []) if accompaniment_data else [],
                            "hortifrutti": accompaniment_data.get("hortifruti", []) if accompaniment_data else [],
                            "cold_cuts": accompaniment_data.get("frios", []) if accompaniment_data else [],
                            "grocery": accompaniment_data.get("mercearia", []) if accompaniment_data else []
                        }
                    })

        table_data = self.create_table(menu_data, menu_name)
        
        self.create_docx(table_data, f'temp/{ menu_name }.docx')

        excel_file_path = f'temp/{menu_name}.xlsx'
        if os.path.exists(excel_file_path):
            os.remove(excel_file_path)

        temp_list = os.listdir('temp')
        files_list = [file for file in temp_list if menu_name in file]
        files_list = [f'temp/{file}' for file in files_list]

        email_response = self.email.send_media_email(subject="Menu", recipients=[token_data.get("email")], attachment_paths=files_list)

        for file in files_list:
            if os.path.exists(file):
                os.remove(file)

        if not email_response:
            raise Exception("Failed to send email")
        
        return { "success": True, "message": "Menu created successfully" }
        
    @MenuDatabase.log
    def create_table(self, menu_data, menu_name):
        writer = pd.ExcelWriter(f'temp/{ menu_name }.xlsx', engine='xlsxwriter')
        monthly_ingredients = {'protein': set(), 'hortifrutti': set(), 'cold_cuts': set(), 'grocery': set()}

        for week_choice, week_data in menu_data.items():
            table_data = []

            for weekday, day_data in week_data.items():

                main_dish = day_data["main_dish"]["dish"]
                salad = day_data["salad"]["dish"]
                accompaniment = ", ".join([accompaniment["dish"] for accompaniment in day_data["accompaniment"]])

                main_dish_ingredients = ", ".join(day_data["main_dish"]["ingredients"]["protein"] + day_data["main_dish"]["ingredients"]["hortifrutti"] + day_data["main_dish"]["ingredients"].get("cold_cuts", []) + day_data["main_dish"]["ingredients"].get("grocery", []))
                salad_ingredients = ", ".join(day_data["salad"]["ingredients"]["protein"] + day_data["salad"]["ingredients"]["hortifrutti"] + day_data["salad"]["ingredients"].get("cold_cuts", []) + day_data["salad"]["ingredients"].get("grocery", []))
                accompaniment_ingredients = ", ".join([", ".join(accompaniment["ingredients"]["protein"] + accompaniment["ingredients"]["hortifrutti"] + accompaniment["ingredients"].get("cold_cuts", []) + accompaniment["ingredients"].get("grocery", [])) for accompaniment in day_data["accompaniment"]])

                table_data.append([self.handle_week_day(weekday), main_dish, salad, accompaniment, main_dish_ingredients, salad_ingredients, accompaniment_ingredients])

                for ingredient_type in ['protein', 'hortifrutti', 'cold_cuts', 'grocery']:
                    monthly_ingredients[ingredient_type].update(
                        day_data["main_dish"]["ingredients"].get(ingredient_type, []) 
                        + day_data["salad"]["ingredients"].get(ingredient_type, []) 
                        + list(set([ingredient for accompaniment in day_data["accompaniment"] for ingredient in accompaniment["ingredients"].get(ingredient_type, [])]))
                    )

            columns = ["Dia da Semana", "Prato Principal", "Salada", "Acompanhamento", "Ingredientes Prato Principal", "Ingredientes Salada", "Ingredientes Acompanhamentos"]
            df = pd.DataFrame(table_data, columns=columns)
            df.to_excel(writer, sheet_name=f"Semana {week_choice}", index=False)

            # Create a sheet for week ingredients
            week_ingredients = {
                "protein": set(),
                "hortifrutti": set(),
                "cold_cuts": set(),
                "grocery": set(),
            }

            self.logger.info("teste")
            for day_data in week_data.values():
                for ingredient_type in ['protein', 'hortifrutti', 'cold_cuts', 'grocery']:
                    week_ingredients[ingredient_type].update(
                        day_data["main_dish"]["ingredients"].get(ingredient_type, [])
                        + day_data["salad"]["ingredients"].get(ingredient_type, [])
                        + list(set([ingredient for accompaniment in day_data["accompaniment"] for ingredient in accompaniment["ingredients"].get(ingredient_type, [])])))

            week_length = max(len(week_ingredients["protein"]), len(week_ingredients["hortifrutti"]), len(week_ingredients["cold_cuts"]), len(week_ingredients["grocery"]))
            week_proteins = list(week_ingredients["protein"]) + [None] * (week_length - len(week_ingredients["protein"]))
            week_hortifruti = list(week_ingredients["hortifrutti"]) + [None] * (week_length - len(week_ingredients["hortifrutti"]))
            week_cold_cuts = list(week_ingredients["cold_cuts"]) + [None] * (week_length - len(week_ingredients["cold_cuts"]))
            week_grocery = list(week_ingredients["grocery"]) + [None] * (week_length - len(week_ingredients["grocery"]))
            
            week_ingredients_data = {
                "Proteina": week_proteins,
                "Hortifruti": week_hortifruti,
                "Frios": week_cold_cuts,
                "Mercearia": week_grocery
            }
            week_ingredients_df = pd.DataFrame(week_ingredients_data, columns=["Proteina", "Hortifruti", "Frios", "Mercearia"])
            week_ingredients_df.to_excel(writer, sheet_name=f"Ingredientes Semana {week_choice}", index=False)

        max_length = max(len(monthly_ingredients["protein"]), len(monthly_ingredients["hortifrutti"]), len(monthly_ingredients["cold_cuts"]), len(monthly_ingredients["grocery"]))

        # Extend the shorter lists with None to match the max_length
        proteins = list(monthly_ingredients["protein"]) + [None] * (max_length - len(monthly_ingredients["protein"]))
        hortifruti = list(monthly_ingredients["hortifrutti"]) + [None] * (max_length - len(monthly_ingredients["hortifrutti"]))
        cold_cuts = list(monthly_ingredients["cold_cuts"]) + [None] * (max_length - len(monthly_ingredients["cold_cuts"]))
        grocery = list(monthly_ingredients["grocery"]) + [None] * (max_length - len(monthly_ingredients["grocery"]))

        monthly_ingredients_df = pd.DataFrame({
            "Proteina": proteins,
            "Hortifruti": hortifruti,
            "Frios": cold_cuts,
            "Mercearia": grocery
        })

        monthly_ingredients_df.to_excel(writer, sheet_name="Ingredientes do Mês", index=False)

        writer.close()

        return pd.read_excel(f'temp/{ menu_name }.xlsx', sheet_name=None)

    @MenuDatabase.log
    def add_table_from_df(self, df, doc):
        table = doc.add_table(rows=(df.shape[0] + 1), cols=df.shape[1])
        table.style = 'Table Grid'
        total_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
        cell_width = total_width / df.shape[1]

        for row in table.rows:
            for cell in row.cells:
                cell.width = cell_width
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.word_wrap = False

        for j, col_name in enumerate(df.columns):
            cell = table.cell(0, j)
            cell.text = col_name
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                cell = table.cell(i + 1, j)
                if pd.notnull(df.iloc[i, j]):
                    cell.text = str(df.iloc[i, j])
                else:
                    cell.text = ''
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    @MenuDatabase.log
    def create_paragraph(self, document, sheet_name, df_week):
        document.add_paragraph().add_run(sheet_name + ":\n").bold = True
        for index, row in df_week.iterrows():
            day_summary = ""
            day_name = row['Dia da Semana']
            if pd.notna(row['Prato Principal']):
                day_summary += f"{row['Prato Principal']}"
            if pd.notna(row['Salada']):
                day_summary += f", {row['Salada']}"
            if pd.notna(row['Acompanhamento']):
                day_summary += f", {row['Acompanhamento']}"
            paragraph = document.add_paragraph()
            paragraph.add_run(day_name + ": ").bold = True
            paragraph.add_run(day_summary)

    @MenuDatabase.log
    def create_docx(self, dataframe: pd.DataFrame, name: str):
        all_sheet_names = dataframe.keys()

        for sheet_name in all_sheet_names:
            if sheet_name.startswith("Ingredientes Semana"):
                continue
            doc = Document()

            header = doc.sections[0].header
            header_image_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.realpath(__file__)))),"public","static","images","logo2.png")

            header_paragraph = header.paragraphs[0]
            run = header_paragraph.add_run()
            run.add_picture(header_image_path, width=Inches(2))
            run.font.size = Pt(12)
            
            df_week = dataframe[sheet_name]
            if sheet_name == "Ingredientes do Mês":
                doc.add_heading('Ingredientes do Mês', level=1)
                self.add_table_from_df(df_week, doc)
            elif sheet_name.startswith("Semana"):
                doc.add_heading('Resumo do Cardápio Semanal', level=1)
                self.create_paragraph(doc, sheet_name, df_week)
                doc.add_page_break()
                df_week = dataframe[f"Ingredientes Semana {sheet_name.split(' ')[-1]}"]
                self.add_table_from_df(df_week, doc)

            doc.save(f"{name} - {sheet_name}.docx")